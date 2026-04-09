#!/usr/bin/env python3
"""Generate GoMate pitch.docx — professional, branded document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Brand colors ──
PULSE_GREEN = RGBColor(0x22, 0xC5, 0x5E)
MIDNIGHT = RGBColor(0x0F, 0x17, 0x2A)
FOREST_DARK = RGBColor(0x1B, 0x3A, 0x2D)
FOREST_MID = RGBColor(0x23, 0x4D, 0x3A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)
LIGHT_GREEN_BG = RGBColor(0xEC, 0xFD, 0xF5)
MUTED_TEXT = RGBColor(0x64, 0x74, 0x8B)
DARK_TEXT = RGBColor(0x1E, 0x29, 0x3B)

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "GoMate_Pitch.docx")
LOGO_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "public", "images", "gomate-logo.png")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_TEXT

# ── Helper functions ──

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_green_line():
    """Add a thin green divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("━" * 60)
    run.font.color.rgb = PULSE_GREEN
    run.font.size = Pt(6)

def add_section_header(text):
    """Add a branded section header with green accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(6)
    # Green bar character
    run = p.add_run("▌ ")
    run.font.color.rgb = PULSE_GREEN
    run.font.size = Pt(18)
    run.bold = True
    # Header text
    run = p.add_run(text.upper())
    run.font.color.rgb = MIDNIGHT
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = 'Calibri'
    add_green_line()

def add_subsection_header(text):
    """Add a subsection header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = FOREST_DARK
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = 'Calibri'

def add_body(text):
    """Add body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p

def add_bold_body(bold_text, normal_text):
    """Add body with bold lead-in."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = MIDNIGHT
    run = p.add_run(normal_text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p

def add_quote(text):
    """Add a styled quote block."""
    # Create a single-cell table as a quote card
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "ECFDF5")
    # Left border green
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:color="22C55E" w:space="0"/>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    for i, line in enumerate(text.split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = FOREST_DARK
        run.italic = True
        run.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_card_table(headers, rows, col_widths=None):
    """Add a branded table with green header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "1B3A2D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.bold = True
        run.font.name = 'Calibri'

    # Style data rows
    for r, row_data in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else "F1F5F9"
        for c, val in enumerate(row_data):
            cell = table.cell(r + 1, c)
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            # First column bold
            run = p.add_run(val)
            if c == 0:
                run.bold = True
                run.font.color.rgb = MIDNIGHT
            else:
                run.font.color.rgb = DARK_TEXT
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:color="E2E8F0" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_bullet(text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = Pt(16)
    # Bullet character
    run = p.add_run("● ")
    run.font.color.rgb = PULSE_GREEN
    run.font.size = Pt(9)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = MIDNIGHT
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT

def add_step_card(number, title, description):
    """Add a numbered step with green number."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"  {number}  ")
    run.font.color.rgb = PULSE_GREEN
    run.font.size = Pt(20)
    run.bold = True
    run = p.add_run(title)
    run.font.color.rgb = MIDNIGHT
    run.font.size = Pt(12)
    run.bold = True
    # Description
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run(description)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT

def add_metric_row(metrics):
    """Add a row of key metrics in card style."""
    table = doc.add_table(rows=2, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (value, label) in enumerate(metrics):
        # Value cell
        cell = table.cell(0, i)
        set_cell_shading(cell, "ECFDF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.color.rgb = PULSE_GREEN
        run.font.size = Pt(22)
        run.bold = True
        run.font.name = 'Calibri'
        # Label cell
        cell = table.cell(1, i)
        set_cell_shading(cell, "ECFDF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.color.rgb = MUTED_TEXT
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

    # Remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:left w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:color="auto" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Spacer
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Logo
if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))

# Spacer
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GoMate")
run.font.size = Pt(42)
run.font.color.rgb = MIDNIGHT
run.bold = True
run.font.name = 'Calibri'

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("The Operating System for Cross-Border Life")
run.font.size = Pt(16)
run.font.color.rgb = PULSE_GREEN
run.font.name = 'Calibri'

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 30)
run.font.color.rgb = PULSE_GREEN
run.font.size = Pt(8)

# Spacer
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Hook quote
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("If you miss a single deadline when moving abroad,")
run.font.size = Pt(13)
run.font.color.rgb = DARK_TEXT
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("you can face fines, delays — or even deportation.")
run.font.size = Pt(13)
run.font.color.rgb = DARK_TEXT
run.italic = True

# Core claim
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Moving abroad isn't an information problem.")
run.font.size = Pt(14)
run.font.color.rgb = MIDNIGHT
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
run = p.add_run("It's an execution problem.")
run.font.size = Pt(14)
run.font.color.rgb = PULSE_GREEN
run.bold = True

# Spacer
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Footer info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Axel Cornelius & Roselle")
run.font.size = Pt(11)
run.font.color.rgb = MUTED_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026")
run.font.size = Pt(11)
run.font.color.rgb = MUTED_TEXT

# Page break
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# THE PROBLEM
# ══════════════════════════════════════════════════════════════

add_section_header("The Problem")

add_body("Moving abroad is one of life's biggest projects — and one of the worst supported.")

add_quote("It's not the information that's missing. It's the execution.")

add_body(
    "Anyone can google a visa in five minutes. But everything else — where to register "
    "your residency, how to open a bank account, what deadlines you have, in what order "
    "things need to happen — it's a mess. And that mess is unique to every country, "
    "every situation, every person."
)

add_bold_body(
    "Time-critical deadlines slip through the cracks. ",
    "Population registration within 7 days. Tax registration within 30 days. "
    "Residence permit before your first day of work. Miss a deadline and you risk "
    "fines, deportation, or months of extra waiting."
)

add_bold_body(
    "No service covers the full journey. ",
    "Relocation agencies cost $2,000–5,000 and focus on corporations, not individuals. "
    "Expat forums have anecdotes, not structured plans. Google gives generic lists that "
    "don't account for your specific situation. ChatGPT hallucinates, has no compliance "
    "tracking, and no way to verify against official sources."
)

add_bold_body(
    "The result: ",
    "Many spend 50–100 hours on research, still miss critical steps, "
    "and feel lost when they arrive."
)

# Key metrics
add_metric_row([
    ("50–100h", "Research time wasted"),
    ("$2–5k", "Agency cost (corporate only)"),
    ("7 days", "Miss one deadline = fines"),
    ("0", "Solutions for individuals"),
])


# ══════════════════════════════════════════════════════════════
# WHY NOW
# ══════════════════════════════════════════════════════════════

add_section_header("Why Now")

add_body("Four forces are colliding — creating a window:")

add_bold_body("1. The remote work explosion. ",
    "Post-COVID, remote work has gone from exception to norm. Millions of people can "
    "now choose where they live — and are doing so. Digital nomad visas have grown "
    "from 5 to 50+ countries in three years.")

add_bold_body("2. Governments are tightening regulations. ",
    "Countries are introducing stricter compliance requirements, shorter deadlines, and "
    "higher fines. Portugal, Spain, Germany — all have tightened their rules in 2024–2026. "
    "The consequences of missing a deadline are growing.")

add_bold_body("3. Individuals are on their own. ",
    "Companies have relocation agencies. Individuals have Google. The gap is enormous — "
    "and it's growing as more people relocate independently.")

add_bold_body("4. AI makes this possible only now. ",
    "Three years ago, this product was impossible. Web research pipelines, LLM synthesis, "
    "and real-time personalization — all of this is functional only now. GoMate isn't an "
    "idea waiting for the technology. The technology has arrived.")


# ══════════════════════════════════════════════════════════════
# ORIGIN
# ══════════════════════════════════════════════════════════════

add_section_header("Origin")

add_body("GoMate shouldn't have needed to be built.")

add_body(
    "My fiancée is from the Philippines. I'm from Sweden. We want to live together, "
    "which means one of us needs to relocate."
)

add_body(
    "I thought it would be straightforward — figure out the visa, apply, done. "
    "But it wasn't. I spent hours jumping between Reddit threads, forums, and "
    "government websites trying to understand what you actually need to do. "
    "Not just the visa, but everything around it — where to start, what happens "
    "when you land, what's expected of you locally."
)

add_body("There was plenty of information. But nothing tailored to our situation.")

add_quote(
    "That's when it clicked: this isn't an information problem —\n"
    "it's an execution problem.\n"
    "And there's no tool that actually helps you get through\n"
    "the entire process step by step."
)

add_body(
    "What can I expect as a Swede in the Philippines? What can she expect as a "
    "Filipina in Sweden? The answers exist — scattered across hundreds of sites, "
    "in different formats, in different languages. Someone needs to pull it together "
    "and make it executable."
)

add_body("So we built it. Not as a guide — but as a system that actually runs the process.")


# ══════════════════════════════════════════════════════════════
# THE SOLUTION
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
add_section_header("The Solution")

add_body(
    "GoMate is an AI-powered relocation operating system that builds a complete, "
    "personalized relocation plan based on a 15-minute chat interview — and then "
    "follows you through the entire settling-in process."
)

add_step_card("01", "Describe your situation (15 min)",
    "An AI assistant asks questions in a natural conversation. No forms, no dropdowns "
    "— just a conversation. Under the surface, GoMate builds a profile with 65+ data "
    "points: citizenship, destination, purpose, family situation, budget, work experience, "
    "language skills, health needs, and more.")

add_step_card("02", "Get your plan (automatically)",
    "Once the profile is complete, an AI-driven research pipeline searches official "
    "sources (embassies, immigration websites, government agencies) via web scraping, "
    "analyzes and synthesizes the information with AI, and builds a tailored guide "
    "with 15+ sections.")

add_step_card("03", "Act with confidence",
    "Your plan contains everything you need:")

add_card_table(
    ["Section", "What you get"],
    [
        ["Visa recommendations", "3–8 visa types ranked by your profile, with requirements, cost, processing time, and step-by-step instructions"],
        ["Budget & cost of living", "Real-time data for your city with month-by-month budget (minimum vs comfortable)"],
        ["Document checklist", "Prioritized documents (critical/high/medium/low) with where to obtain them"],
        ["Housing guide", "Rent by area, platforms, deposit rules"],
        ["Healthcare, banking, jobs, culture", "Everything tailored to your country, your city, your situation"],
        ["Flight search", "Comparisons from 5 booking sites"],
        ["Timeline", "Milestones from now to arrival"],
    ]
)

add_step_card("04", "Land and get established (Pro+)",
    "When you arrive, you activate post-arrival mode:")

add_bullet(" A bank account can't be opened until you've registered your residency — and GoMate knows that", "Task graph with dependencies —")
add_bullet(' "Register with the tax authority within 7 days"', "Deadlines tied to your arrival date — ")
add_bullet(" Overdue / urgent / approaching — so you never miss anything legally critical", "Compliance alerts —")
add_bullet(" Ask anything about your new city and get answers based on your profile and plan", "AI chat for questions —")


# ══════════════════════════════════════════════════════════════
# WEDGE
# ══════════════════════════════════════════════════════════════

add_section_header("Wedge")

add_body('We don\'t start with "everyone who moves abroad." We start here:')

add_quote("Professionals and couples relocating within and to/from Europe.")

add_body(
    "This is the segment with the highest frequency, highest compliance complexity, "
    "and highest urgency — and where there is currently no solution for individuals."
)

add_body(
    "EU freedom of movement sounds simple, but the reality is different: every country "
    "has unique requirements for residency registration, tax registration, insurance, "
    "and bank account opening. Deadlines are strict. Consequences are real."
)

add_bullet(" High relocation frequency, short cycles", "Fast feedback loops —")
add_bullet(" The same compliance requirements recur country by country", "Repeatable flows —")
add_bullet(" Deadlines with legal consequences drive willingness to pay", "High urgency —")

add_body("From there, we expand to digital nomads, global relocations, and more country combinations.")


# ══════════════════════════════════════════════════════════════
# TRUST
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
add_section_header("Trust")

add_quote("We don't generate advice.\nWe generate verifiable, time-bound instructions\ntied to official sources.")

add_body("Every recommendation in GoMate is:")

add_bullet(" Embassies, immigration authorities, government agencies. Linked so the user can verify.", "Backed by an official source —")
add_bullet(" The user sees when the information was retrieved.", "Timestamped —")
add_bullet(' Full (complete data from official sources), partial (incomplete), or fallback (AI-based). Nothing is presented as more certain than it is.', "Quality-rated —")

add_body(
    "When a profile changes or research is older than 7 days, the guide is automatically "
    "marked as stale. Nothing is ever presented as current unless it is."
)

add_body(
    "Next step is continuous re-validation: when a source changes, your plan updates "
    "— before it becomes your problem."
)

add_body(
    "Profile data extracted from the chat interview is tagged with a confidence level: "
    "explicit, inferred, or assumed. Assumed data is always confirmed with the user "
    "before it's used."
)

add_body(
    "GoMate is not a lawyer, advisor, or government authority. We help you find, structure, "
    "and execute — but the final responsibility is always yours. This is communicated "
    "clearly in the product."
)


# ══════════════════════════════════════════════════════════════
# MOAT
# ══════════════════════════════════════════════════════════════

add_section_header("Moat")

add_quote("Most players in this space build guides.\nWe build infrastructure.")

add_bold_body("1. The compliance engine. ",
    "A DAG-based task graph where every task has dependencies, deadlines relative to "
    "arrival date, legal-requirement flags, and urgency computation. It's not a checklist "
    "— it's an operational system that computes what you need to do, in what order, and "
    "warns you when you're about to miss something. Building this correctly took 11 build "
    "phases. It's not a feature — it's an engine.")

add_bold_body("2. Profile depth. ",
    "65+ data points with intelligent branching — your purpose (work/study/digital nomad/"
    "family) determines which questions are asked, which research is triggered, and which "
    "output you get. Every new profile makes the platform's output quality better.")

add_bold_body("3. Research pipeline with web grounding. ",
    "Not just LLM generation — but real-time scraping of official sources + AI synthesis. "
    "Every result has source attribution. Next step is continuous monitoring: when a source "
    "changes, it gets flagged.")

add_quote(
    "This is not a feature problem. It's a systems problem.\n"
    "And systems — once built correctly —\n"
    "are extremely hard to replace."
)

add_body(
    "Someone can build a relocation chatbot in a weekend. No one can build a compliance "
    "engine with DAG dependencies, deadline computation, web-grounded research, and "
    "65-field profile personalization in a weekend."
)


# ══════════════════════════════════════════════════════════════
# BUSINESS MODEL
# ══════════════════════════════════════════════════════════════

add_section_header("Business Model")

add_card_table(
    ["Tier", "Price", "Includes"],
    [
        ["Free", "$0", "Chat interview + profile + overview"],
        ["Pro", "$189 (one-time)", "Full guide, visa research, budget, documents, flight search"],
        ["Pro+", "$39/mo", "Everything in Pro + unlimited plans, post-arrival compliance OS, AI assistant"],
    ]
)

add_body("Pro+ bundles: 3 months ($95), 6 months ($169), annual ($289).")

add_bold_body("Why this pricing: ",
    "A relocation agency costs $2,000–5,000. GoMate delivers 80% of the value at 5% of "
    "the price. $189 for Pro is less than a single hour of consultation — but you get "
    "a complete, AI-driven plan.")

add_bold_body("Unit economics are strong: ",
    "Marginal cost per user is LLM calls (~$0.50–1.50) + Firecrawl search (~$0.20–0.50). "
    "No human advisors, no manual research. Gross margin >90%.")

add_metric_row([
    (">90%", "Gross margin"),
    ("~$1–2", "Cost per plan"),
    ("$189+", "Revenue per user"),
    ("95x+", "LTV/CAC potential"),
])


# ══════════════════════════════════════════════════════════════
# DISTRIBUTION
# ══════════════════════════════════════════════════════════════

add_section_header("Distribution")

add_subsection_header("Growth Engine")

add_body(
    "Every relocation plan we generate can become a structured, indexed guide. Over time, "
    "this creates thousands of high-intent landing pages — continuously updated with real "
    "user patterns — where every guide improves the next."
)

add_quote("This is not content.\nIt's a self-improving distribution engine.")

add_subsection_header("Phase 1: Organic + SEO (0–6 months)")

add_bold_body("Country-specific guides as SEO magnets. ",
    'GoMate already generates 15+ section guides for every country-destination combination. '
    'Publish the best as standalone landing pages: "Moving to Spain from the UK — complete '
    'guide 2026." Every guide is a top-of-funnel leading to signup.')

add_bold_body("Expat communities. ",
    "Reddit (r/expats, r/digitalnomad, r/IWantOut), Facebook groups, InterNations. "
    "Not spam — genuine presence with answers that demonstrate the product's depth.")

add_bold_body("Content on TikTok/YouTube. ",
    'Short videos: "3 deadlines you MUST NOT miss when moving to Germany." '
    "Relocation content has high engagement and low competition.")

add_subsection_header("Phase 2: Partnerships (6–12 months)")

add_bold_body("Expat platforms. ",
    'Integration with InterNations, Expatica, or similar — "Powered by GoMate compliance tracking."')

add_bold_body("Companies with international hires. ",
    "Small businesses that can't afford relocation agencies but hire internationally. "
    "GoMate as the onboarding tool for new hires.")

add_bold_body("Government agencies and educational institutions. ",
    'Universities with international students. Migration organizations. GoMate as a "recommended tool."')

add_subsection_header("Phase 3: Platform Expansion (12+ months)")

add_body(
    "Every completed relocation generates data about what works in which countries. "
    "This data makes the next user's plan better — a data flywheel that accelerates over time."
)


# ══════════════════════════════════════════════════════════════
# MARKET
# ══════════════════════════════════════════════════════════════

doc.add_page_break()
add_section_header("Market")

add_metric_row([
    ("3M+", "Relocations/year (OECD)"),
    ("$20B+", "Global relocation market"),
    ("50+", "Countries with nomad visas"),
    ("20%+", "Annual nomad market growth"),
])

add_card_table(
    ["Alternative", "Problem"],
    [
        ["Relocation agencies", "$2,000–5,000, focused on corporations"],
        ["Expat forums", "Anecdotal, unstructured, impersonal"],
        ["Generic guides", "Not personalized, quickly outdated"],
        ["ChatGPT directly", "Hallucinates, no compliance tracking, no web grounding"],
    ]
)

add_body(
    "GoMate is the first service that gives individuals access to personalized relocation "
    "intelligence with compliance tracking — at a fraction of the cost."
)


# ══════════════════════════════════════════════════════════════
# TEAM
# ══════════════════════════════════════════════════════════════

add_section_header("Team")

add_bold_body("Axel — Product, System Design & AI Engineering", "")
add_body(
    "Not a programmer — a problem solver. Understands systems, how they're expected to work, "
    "and what the output should look like. Finds abstract solutions to complex problems. "
    "Built all of GoMate — architecture, compliance engine, research pipeline, 24 system "
    "definitions, engineering contracts — using AI as the tool. Uses ChatGPT for system "
    "design and Claude Code for implementation, governed through detailed contracts and "
    "specifications. Background in scaling GTM systems, SEO/AEO infrastructure, and "
    "onboarding processes for sales and marketing teams."
)

add_bold_body("Roselle — Design & User Experience", "")
add_body(
    "Trained teacher with a background as Virtual Assistant and Executive Assistant. "
    "Designs all of GoMate's user flows — taking a chaotic problem and making it navigable. "
    "The teaching background gives the ability to take something complex and make it "
    "understandable. Operational experience from the EA role provides deep understanding "
    "of complex processes, deadlines, and workflows — exactly what GoMate's UX demands."
)

add_quote(
    "We didn't start with a market thesis.\n"
    "We started with a real cross-border relocation\n"
    "between Sweden and the Philippines —\n"
    "and built the tool we couldn't find."
)


# ══════════════════════════════════════════════════════════════
# TECHNOLOGY
# ══════════════════════════════════════════════════════════════

add_section_header("Technology")

add_card_table(
    ["Layer", "Technology"],
    [
        ["Frontend", "Next.js 16, React 19, Tailwind CSS, shadcn/ui"],
        ["Backend", "Vercel serverless, Supabase (PostgreSQL + auth)"],
        ["AI — chat & extraction", "GPT-4o via OpenRouter"],
        ["AI — generation & research", "Claude Sonnet 4 via OpenRouter"],
        ["Web research", "Firecrawl (search + scrape)"],
        ["Data", "Numbeo (cost of living), 5 flight booking sites"],
    ]
)


# ══════════════════════════════════════════════════════════════
# STATUS
# ══════════════════════════════════════════════════════════════

add_section_header("Status")

add_body("GoMate v1 is live. 11 build phases completed. Core functionality works end-to-end:")

add_bullet("Chat interview with profile building (65+ fields, intelligent branching)")
add_bullet("AI-driven visa and requirements research (web-grounded, source attribution)")
add_bullet("Complete guide generation (15+ sections, staleness detection)")
add_bullet("Cost of living analysis with currency conversion")
add_bullet("Document checklist with prioritization")
add_bullet("Post-arrival task graph with DAG dependencies and deadline tracking")
add_bullet("Compliance alerts (overdue/urgent/approaching)")
add_bullet("Confidence scoring on extracted profile data")
add_bullet("Subscription management (Free / Pro / Pro+)")


# ══════════════════════════════════════════════════════════════
# VISION
# ══════════════════════════════════════════════════════════════

add_section_header("Vision")

add_quote(
    "GoMate is not a relocation tool.\n"
    "It is the first operating system for cross-border life —\n"
    "starting with relocation, expanding into everything\n"
    "that happens after."
)

add_card_table(
    ["Phase", "Expansion"],
    [
        ["v1 (now)", "Relocation OS — visa, guide, compliance, post-arrival"],
        ["v2", "Job system — matching user profiles to open positions in the destination country"],
        ["v3", "Artifact generation — CV adaptation, cover letters, application materials"],
        ["v4", "Partner integration — direct connections to housing platforms, banks, insurance providers"],
        ["v5", "Community — matching with others moving to the same city"],
        ["Long-term", "Tax, insurance, pensions, citizenship — everything that moves across borders"],
    ]
)

add_body("Every step adds more data, more value creation, and higher switching cost.")

# Spacer
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Closing line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
# Green bar
run = p.add_run("━" * 20)
run.font.color.rgb = PULSE_GREEN
run.font.size = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("GoMate")
run.font.size = Pt(20)
run.font.color.rgb = MIDNIGHT
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The operating system for cross-border life.")
run.font.size = Pt(13)
run.font.color.rgb = PULSE_GREEN
run.italic = True


# ── Save ──
doc.save(OUTPUT_PATH)
print(f"✓ Saved to {OUTPUT_PATH}")
